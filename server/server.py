from docx import Document

# Создаем новый документ
doc = Document()

# Заголовок документа
doc.add_heading('Заявка на бронирование квартиры', level=1)

# Информация о заявителе
doc.add_heading('Информация о заявителе', level=2)
doc.add_paragraph('ФИО заявителя: _______________________')
doc.add_paragraph('Контактный телефон: _______________________')
doc.add_paragraph('Электронная почта: _______________________')
doc.add_paragraph('Адрес постоянного проживания: _______________________')
doc.add_paragraph('Дата подачи заявки: _______________________')

# Информация о квартире
doc.add_heading('Информация о квартире', level=2)
doc.add_paragraph('Номер квартиры: _______________________')
doc.add_paragraph('Этаж: _______________________')
doc.add_paragraph('Площадь квартиры (м²): _______________________')
doc.add_paragraph('Цена квартиры: _______________________')
doc.add_paragraph('Описание квартиры:')
doc.add_paragraph('___________________________________________________________________________')

# Дата бронирования
doc.add_heading('Дата бронирования', level=2)
doc.add_paragraph('Предполагаемая дата бронирования: _______________________')

# Дополнительные условия
doc.add_heading('Дополнительные условия и комментарии', level=2)
doc.add_paragraph('___________________________________________________________________________')
doc.add_paragraph('___________________________________________________________________________')

# Подпись и дата
doc.add_heading('Подпись и дата', level=2)
doc.add_paragraph('Подпись заявителя: _______________________')
doc.add_paragraph('Дата: _______________________')

# Сохраняем документ
doc_path = r"C:\Users\varla\Downloads\Заявка_на_бронирование_квартиры.docx"
doc.save(doc_path)

doc_path
